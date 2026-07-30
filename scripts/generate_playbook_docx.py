#!/usr/bin/env python3
"""Generate the team's Google Ads Playbook docx from google-ads.md.

Replicates the original hand-made docx design system (Calibri, Tailwind
palette, cover cards, dark section banners, zebra tables) while rendering
the body straight from the markdown source of truth
(.claude/stacks/distribution/google-ads.md) — so the docx can never drift
behind the playbook again. The cover cards / at-a-glance numbers are the
only hardcoded content (see CARDS/GLANCE below); keep them in sync when
the standardized Phase numbers change.

Entry point: `make playbook-docx` (manages a local .venv-docx with
python-docx; pass OUT=<path> to override the output location).

Direct usage:
    python scripts/generate_playbook_docx.py --repo . --out out.docx
"""

import argparse
import datetime
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor

# ---------------------------------------------------------------- palette
INK = "111827"        # gray-900
BODY = "374151"       # gray-700
MUTED = "6B7280"      # gray-500
BANNER_BG = "1F2937"  # gray-800
BANNER_FG = "F9FAFB"
ZEBRA = "F9FAFB"
RULE = "D1D5DB"       # gray-300 table borders
CODE_BG = "F3F4F6"    # gray-100
QUOTE_BG = "EFF6FF"   # blue-50
QUOTE_BAR = "1D4ED8"
AMBER_BG = "FEF3C7"
AMBER_FG = "92400E"

CARDS = [
    ("PHASE 1 — DEMAND", "$250 cap", "100 paid clicks · 7 days · $20/day", "ECFEFF", "0E7490"),
    ("PHASE 2 — VALUE", "$750 cap", "300 paid clicks · 21 days · $20/day", "D1FAE5", "065F46"),
    ("BIDDING (MANUAL CPC)", "$1.40 → $2.50", "start bid → ceiling, never bid above", "F3F4F6", INK),
    ("PRICED > $50/MO?", "Ask your Team Lead", "approved raise — click targets never change", "FEF3C7", AMBER_FG),
]

GLANCE = [
    ("", "PHASE 1 — DEMAND SCREEN", "PHASE 2 — VALUE SCREEN"),
    ("Question answered", "Do people want it? (signups)", "Will they pay? (pay-intent)"),
    ("Click target (stop A)", "100 paid clicks", "300 paid clicks"),
    ("Budget MAX cap (stop B)", "$250", "$750"),
    ("Daily budget", "$20/day", "$20/day"),
    ("Duration / End date", "7 days", "21 days"),
    ("Starting bid", "$1.40", "$1.40 (or Phase 1's actual avg CPC)"),
    ("CPC ceiling", "$2.50", "$2.50"),
    ("Daily check", "`/iterate --check`", "`/iterate --check`"),
    ("Verdict (Team Lead runs)", "`/iterate --cross`", "`/iterate --cross --phase2`"),
]

NOTE_BOX = (
    "Priced > $50/mo and $2.50 can't buy volume? Ask your Team Lead for an approved raise — "
    "the approved ceiling/cap/daily replace the $ numbers above for that MVP; click targets "
    "never change (see Appendix — Cost Model)."
)

# md `## Heading` -> docx banner title (order = document order)
SECTIONS = [
    ("Phase 1 Playbook", "Phase 1 Playbook"),
    ("Phase 2 Playbook (Value Screen)", "Phase 2 Playbook (Value Screen)"),
    ("Cost Model", "Appendix — Cost Model (lead-facing economics)"),
]

MONO = "Consolas"


# ---------------------------------------------------------------- xml helpers
def _shade(el_pr, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    el_pr.append(shd)


def shade_cell(cell, fill):
    _shade(cell._tc.get_or_add_tcPr(), fill)


def shade_run(run, fill):
    _shade(run._r.get_or_add_rPr(), fill)


def set_table_borders(table, color=RULE, size=4):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)
        borders.append(el)
    tbl_pr.append(borders)


def no_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tbl_pr.append(borders)


def cell_margins(table, top=60, bottom=60, left=100, right=100):
    tbl_pr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for edge, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tbl_pr.append(mar)


def left_bar(cell, color=QUOTE_BAR, size=18):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    el = OxmlElement("w:left")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size))
    el.set(qn("w:color"), color)
    borders.append(el)
    tc_pr.append(borders)


# ---------------------------------------------------------------- inline md
INLINE_RE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def add_runs(par, text, size=9.5, color=BODY, bold=False):
    """Render **bold** and `code` inline markdown into runs."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            inner = part[2:-2]
            for sub in INLINE_RE.split(inner):
                if not sub:
                    continue
                if sub.startswith("`") and sub.endswith("`"):
                    r = par.add_run(sub[1:-1])
                    r.font.name = MONO
                    shade_run(r, CODE_BG)
                else:
                    r = par.add_run(sub)
                r.font.bold = True
                r.font.size = Pt(size)
                r.font.color.rgb = RGBColor.from_string(INK)
        elif part.startswith("`") and part.endswith("`"):
            r = par.add_run(part[1:-1])
            r.font.name = MONO
            r.font.bold = bold
            r.font.size = Pt(size)
            r.font.color.rgb = RGBColor.from_string(INK)
            shade_run(r, CODE_BG)
        else:
            r = par.add_run(part)
            r.font.bold = bold
            r.font.size = Pt(size)
            r.font.color.rgb = RGBColor.from_string(color)


def para(doc, before=0, after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p


# ---------------------------------------------------------------- md blocks
def parse_blocks(lines):
    """Parse a flat markdown region into (kind, payload) blocks."""
    blocks = []
    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        if not line.strip():
            i += 1
            continue
        if line.startswith("### "):
            blocks.append(("h3", line[4:].strip()))
            i += 1
        elif line.startswith("```"):
            code = []
            i += 1
            while i < n and not lines[i].startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            blocks.append(("code", code))
        elif line.startswith("|"):
            rows = []
            while i < n and lines[i].startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-+:?", c) for c in cells):
                    rows.append(cells)
                i += 1
            blocks.append(("table", rows))
        elif line.startswith("> "):
            quote = []
            while i < n and lines[i].startswith(">"):
                quote.append(lines[i].lstrip(">").strip())
                i += 1
            blocks.append(("quote", " ".join(q for q in quote if q)))
        elif re.match(r"^\d+\. ", line):
            m = re.match(r"^(\d+)\. (.*)$", line)
            num, text = m.group(1), [m.group(2)]
            i += 1
            nested = []
            while i < n:
                nxt = lines[i]
                if re.match(r"^\d+\. ", nxt) or (nxt and not nxt.startswith(" ") and nxt.strip()):
                    break
                if not nxt.strip():
                    # blank: item continues only if the following line is indented
                    if i + 1 < n and lines[i + 1].startswith("    "):
                        nested.append("")
                        i += 1
                        continue
                    break
                if nxt.startswith("    "):
                    nested.append(nxt[4:])
                elif nxt.startswith("   "):
                    text.append(nxt.strip())
                else:
                    break
                i += 1
            blocks.append(("ol", (num, " ".join(text), parse_blocks(nested) if any(s.strip() for s in nested) else [])))
        elif line.startswith("- "):
            items = []
            while i < n and lines[i].startswith("- "):
                item = [lines[i][2:]]
                i += 1
                while i < n and lines[i].startswith("  ") and not lines[i].startswith("- "):
                    item.append(lines[i].strip())
                    i += 1
                items.append(" ".join(item))
            blocks.append(("ul", items))
        else:
            text = [line.strip()]
            i += 1
            while i < n and lines[i].strip() and not re.match(
                r"^(### |```|\||> |- |\d+\. )", lines[i]
            ):
                text.append(lines[i].strip())
                i += 1
            blocks.append(("p", " ".join(text)))
    return blocks


# ---------------------------------------------------------------- renderers
def render_code(doc, code_lines, indent=None):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_table_borders(t)
    cell_margins(t, 80, 80, 120, 120)
    cell = t.rows[0].cells[0]
    shade_cell(cell, CODE_BG)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for j, ln in enumerate(code_lines):
        r = p.add_run(ln if ln else " ")
        r.font.name = MONO
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(INK)
        if j < len(code_lines) - 1:
            r.add_break()
    spacer = para(doc, after=4)
    spacer.paragraph_format.space_before = Pt(0)


def render_table(doc, rows):
    ncols = max(len(r) for r in rows)
    t = doc.add_table(rows=len(rows), cols=ncols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t)
    cell_margins(t, 50, 50, 90, 90)
    for ri, row in enumerate(rows):
        for ci in range(ncols):
            cell = t.rows[ri].cells[ci]
            text = row[ci] if ci < len(row) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if ri == 0:
                shade_cell(cell, BANNER_BG)
                add_runs(p, text, size=8.5, color=BANNER_FG, bold=True)
                for r in p.runs:
                    r.font.color.rgb = RGBColor.from_string(BANNER_FG)
            else:
                if ri % 2 == 0:
                    shade_cell(cell, ZEBRA)
                add_runs(p, text, size=8.5, bold=(ci == 0))
    spacer = para(doc, after=4)
    spacer.paragraph_format.space_before = Pt(0)


def render_quote(doc, text):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_table_borders(t)
    cell_margins(t, 70, 70, 140, 120)
    cell = t.rows[0].cells[0]
    shade_cell(cell, QUOTE_BG)
    left_bar(cell)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_runs(p, text, size=9)
    spacer = para(doc, after=4)
    spacer.paragraph_format.space_before = Pt(0)


def render_blocks(doc, blocks, indent=0.0):
    for kind, payload in blocks:
        if kind == "h3":
            p = para(doc, before=10, after=4)
            add_runs(p, payload, size=12.5, color=INK, bold=True)
            for r in p.runs:
                r.font.color.rgb = RGBColor.from_string(INK)
                r.font.bold = True
        elif kind == "p":
            p = para(doc)
            if indent:
                p.paragraph_format.left_indent = Inches(indent)
            add_runs(p, payload)
        elif kind == "quote":
            render_quote(doc, payload)
        elif kind == "code":
            render_code(doc, payload)
        elif kind == "table":
            render_table(doc, payload)
        elif kind == "ul":
            for item in payload:
                p = para(doc, after=2)
                p.paragraph_format.left_indent = Inches(0.22 + indent)
                r = p.add_run("•  ")
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor.from_string(MUTED)
                add_runs(p, item)
        elif kind == "ol":
            num, text, nested = payload
            p = para(doc, after=2)
            p.paragraph_format.left_indent = Inches(0.22 + indent)
            r = p.add_run(f"{num}.  ")
            r.font.size = Pt(9.5)
            r.font.bold = True
            r.font.color.rgb = RGBColor.from_string(INK)
            add_runs(p, text.replace("[ ]", "☐"))
            if nested:
                render_blocks(doc, nested, indent=indent + 0.28)


def banner(doc, title):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_table_borders(t)
    cell_margins(t, 90, 90, 140, 140)
    cell = t.rows[0].cells[0]
    shade_cell(cell, BANNER_BG)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(title)
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(BANNER_FG)
    spacer = para(doc, after=6)
    spacer.paragraph_format.space_before = Pt(0)


# ---------------------------------------------------------------- md slicing
def slice_section(md_text, heading):
    lines = md_text.splitlines()
    start = None
    for i, ln in enumerate(lines):
        if ln.strip() == f"## {heading}":
            start = i + 1
            break
    if start is None:
        raise SystemExit(f"section '## {heading}' not found in source md")
    end = len(lines)
    for j in range(start, len(lines)):
        if lines[j].startswith("## "):
            end = j
            break
    return lines[start:end]


# ---------------------------------------------------------------- cover
def build_cover(doc, provenance):
    p = para(doc, after=2)
    add_runs(p, "Google Ads Playbook", size=24, color=INK, bold=True)
    p = para(doc, after=2)
    add_runs(p, "Phase 1 (Demand Screen)  ·  Phase 2 (Value Screen)", size=13, color=BODY, bold=True)
    p = para(doc, after=8)
    add_runs(p, provenance, size=8.5, color=MUTED)

    t = doc.add_table(rows=1, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_table_borders(t)
    cell_margins(t, 90, 90, 110, 110)
    for ci, (label, big, desc, fill, fg) in enumerate(CARDS):
        cell = t.rows[0].cells[ci]
        shade_cell(cell, fill)
        p1 = cell.paragraphs[0]
        p1.paragraph_format.space_after = Pt(1)
        r = p1.add_run(label)
        r.font.size = Pt(7.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(fg)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(1)
        r = p2.add_run(big)
        r.font.size = Pt(15)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(fg)
        p3 = cell.add_paragraph()
        p3.paragraph_format.space_after = Pt(0)
        r = p3.add_run(desc)
        r.font.size = Pt(7.5)
        r.font.color.rgb = RGBColor.from_string(BODY)
    para(doc, after=6)

    p = para(doc, after=4)
    add_runs(p, "The numbers at a glance", size=12.5, color=INK, bold=True)
    render_table(doc, [list(r) for r in GLANCE])

    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    no_table_borders(t)
    cell_margins(t, 80, 80, 130, 130)
    cell = t.rows[0].cells[0]
    shade_cell(cell, AMBER_BG)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_runs(p, NOTE_BOX, size=9, color=AMBER_FG)
    para(doc, after=8)


def build_contents(doc, toc):
    p = para(doc, before=4, after=4)
    add_runs(p, "Contents", size=12.5, color=INK, bold=True)
    for level, title in toc:
        p = para(doc, after=1)
        p.paragraph_format.left_indent = Inches(0.12 if level == 2 else 0.38)
        marker = "•  " if level == 2 else "–  "
        r = p.add_run(marker)
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor.from_string(MUTED)
        add_runs(p, re.sub(r"`([^`]*)`", r"\1", title), size=9.5)
    para(doc, after=8)


# ---------------------------------------------------------------- main
def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--repo", default=".")
    ap.add_argument("--out", required=True)
    args = ap.parse_args()

    repo = Path(args.repo)
    md_path = repo / ".claude/stacks/distribution/google-ads.md"
    md_text = md_path.read_text(encoding="utf-8")

    sha = subprocess.run(
        ["git", "-C", str(repo), "rev-parse", "--short", "HEAD"],
        capture_output=True, text=True,
    ).stdout.strip() or "unknown"
    today = datetime.date.today().isoformat()
    provenance = (
        f"Team playbook · {today} · source: .claude/stacks/distribution/google-ads.md "
        f"@ mvp-template {sha}"
    )

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9.5)
    style.font.color.rgb = RGBColor.from_string(BODY)
    sec = doc.sections[0]
    sec.page_width = Emu(7560310)
    sec.page_height = Emu(10692130)
    for m in ("left_margin", "right_margin"):
        setattr(sec, m, Emu(612140))
    sec.top_margin = Emu(540000)
    sec.bottom_margin = Emu(540000)

    sections = [(banner_title, slice_section(md_text, md_heading))
                for md_heading, banner_title in SECTIONS]

    toc = []
    for banner_title, lines in sections:
        toc.append((2, banner_title))
        for ln in lines:
            if ln.startswith("### "):
                toc.append((3, ln[4:].strip()))

    build_cover(doc, provenance)
    build_contents(doc, toc)

    for banner_title, lines in sections:
        banner(doc, banner_title)
        render_blocks(doc, parse_blocks(lines))

    doc.save(args.out)
    print(f"wrote {args.out}")


if __name__ == "__main__":
    main()
