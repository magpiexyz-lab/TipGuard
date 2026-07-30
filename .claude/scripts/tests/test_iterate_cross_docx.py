#!/usr/bin/env python3
"""Tests for .claude/scripts/lib/iterate_cross_docx.py (state x4 decision report).

Run:
  python3 -m pytest .claude/scripts/tests/test_iterate_cross_docx.py -v
  # OR (no pytest dependency):
  python3 .claude/scripts/tests/test_iterate_cross_docx.py
"""
from __future__ import annotations

import os
import sys
import tempfile
from unittest.mock import patch

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "lib"))

import iterate_cross_docx as D  # noqa: E402

TH = {"visitors_floor": 100, "conv_rate_go": 0.06, "max_cpc": 2.5}


def _score(name, verdict, **kw):
    metrics = kw.pop("metrics", {})
    s = {"name": name, "headline_verdict": verdict, "metrics": metrics}
    s.update(kw)
    return s


# ---------- bucket_scores ----------

def test_bucket_basic_partition():
    scores = [
        _score("a", "GO", metrics={"db_signups_real": 5}),
        _score("b", "NO_GO", metrics={"ga_clicks": 100}),
        _score("c", "INSUFFICIENT_DATA"),
        _score("d", "MISSING_PROJECT_NAME"),
        _score("e", "GA_NO_PH_TRACKING"),
        _score("f", "NO_DATA"),
    ]
    b = D.bucket_scores(scores)
    assert [s["name"] for s in b["go"]] == ["a"]
    assert [s["name"] for s in b["no_go"]] == ["b"]
    assert [s["name"] for s in b["insuf"]] == ["c"]
    assert {s["name"] for s in b["fix"]} == {"d", "e", "f"}


def test_bucket_killed_go_becomes_no_go():
    # A GO whose lifecycle is killed must land in NO_GO, not GO, and only once.
    scores = [_score("z", "GO", lifecycle_status="killed", metrics={"db_signups_real": 9, "ga_clicks": 50})]
    b = D.bucket_scores(scores)
    assert [s["name"] for s in b["no_go"]] == ["z"]
    assert b["go"] == []


def test_bucket_money_leak_to_no_go():
    scores = [_score("ml", "INSUFFICIENT_DATA", metrics={"money_leak": True, "ga_clicks": 10})]
    b = D.bucket_scores(scores)
    assert [s["name"] for s in b["no_go"]] == ["ml"]
    assert b["insuf"] == []


def test_bucket_no_double_count():
    # Orphan that is also NO_DATA-shaped lands once (FIX precedence).
    scores = [_score("__orphan_x__", "MISSING_PROJECT_NAME", lifecycle_status="killed")]
    b = D.bucket_scores(scores)
    total = len(b["go"]) + len(b["no_go"]) + len(b["insuf"]) + len(b["fix"]) + len(b["promoted"])
    assert total == 1
    assert [s["name"] for s in b["fix"]] == ["__orphan_x__"]


# ---------- promoted bucket ----------

def test_bucket_promoted_go_lands_in_promoted():
    scores = [_score("h", "GO", lifecycle_status="promoted",
                     metrics={"db_signups_real": 12, "ga_clicks": 371, "ga_clicks_phase2": 258})]
    b = D.bucket_scores(scores)
    assert [s["name"] for s in b["promoted"]] == ["h"]
    assert b["go"] == []


def test_bucket_promoted_no_go_reference_verdict_lands_in_promoted():
    # A promoted MVP whose phase-1 reference verdict is NO_GO (phase1 flight
    # wound down) belongs in the promoted table, NOT the stop list.
    scores = [_score("h", "NO_GO", lifecycle_status="promoted",
                     metrics={"ga_clicks": 300, "ga_clicks_phase2": 300})]
    b = D.bucket_scores(scores)
    assert [s["name"] for s in b["promoted"]] == ["h"]
    assert b["no_go"] == []


def test_bucket_promoted_dead_backend_beats_promoted():
    killed = _score("k", "GO", lifecycle_status="killed", metrics={"ga_clicks": 10})
    leak = _score("l", "INSUFFICIENT_DATA", lifecycle_status="promoted",
                  metrics={"money_leak": True, "ga_clicks": 10})
    deleted = _score("d", "NO_GO", lifecycle_status="promoted",
                     db_unmapped_reason="project_deleted", metrics={"ga_clicks": 10})
    b = D.bucket_scores([killed, leak, deleted])
    assert {s["name"] for s in b["no_go"]} == {"k", "l", "d"}
    assert b["promoted"] == []


def test_bucket_promoted_fix_verdict_stays_in_fix():
    scores = [_score("pf", "GA_NO_PH_TRACKING", lifecycle_status="promoted",
                     metrics={"ga_clicks": 30})]
    b = D.bucket_scores(scores)
    assert [s["name"] for s in b["fix"]] == ["pf"]
    assert b["promoted"] == []


def test_summary_counts_include_promoted():
    scores = [
        _score("a", "GO", metrics={"db_signups_real": 5}),
        _score("h", "GO", lifecycle_status="promoted", metrics={"ga_clicks_phase2": 64}),
    ]
    counts = D.verdict_summary_counts(scores)
    assert counts["go"] == 1
    assert counts["promoted"] == 1


def test_promoted_partition_parity_with_verdicts_module():
    # docx deliberately re-implements the promoted predicate locally (its header
    # forbids importing the heavy verdicts module). This parity matrix pins the
    # two implementations together — any drift fails here.
    from iterate_cross_verdicts import _score_is_promoted as V_pred

    fixtures = []
    for verdict in ("GO", "NO_GO", "INSUFFICIENT_DATA",
                    "MISSING_PROJECT_NAME", "GA_NO_PH_TRACKING", "NO_DATA"):
        for lifecycle in ("active", "promoted", "killed"):
            for extra in (
                {},
                {"metrics": {"money_leak": True}},
                {"db_unmapped_reason": "project_deleted"},
            ):
                s = _score(f"m-{verdict}-{lifecycle}", verdict,
                           lifecycle_status=lifecycle, **{k: v for k, v in extra.items() if k != "metrics"})
                if "metrics" in extra:
                    s["metrics"] = dict(extra["metrics"])
                fixtures.append(s)
    for s in fixtures:
        assert D.is_promoted(s) == V_pred(s), f"parity mismatch on {s['name']}: {s}"


def test_bucket_go_sorted_by_db_signups():
    scores = [
        _score("low", "GO", metrics={"db_signups_real": 2}),
        _score("high", "GO", metrics={"db_signups_real": 20}),
    ]
    b = D.bucket_scores(scores)
    assert [s["name"] for s in b["go"]] == ["high", "low"]


# ---------- mvp_cell_markers ----------

def test_marker_db_deleted_from_unmapped_reason():
    assert "⚠ DB deleted" in D.mvp_cell_markers(_score("x", "NO_GO", db_unmapped_reason="project_deleted"))


def test_marker_db_deleted_from_killed():
    assert "⚠ DB deleted" in D.mvp_cell_markers(_score("x", "NO_GO", lifecycle_status="killed"))


def test_marker_no_db_access():
    mk = D.mvp_cell_markers(_score("x", "NO_GO", db_unmapped_reason="no_match"))
    assert "⚠ no DB access" in mk


def test_marker_partial_tracking():
    mk = D.mvp_cell_markers(_score("x", "NO_GO", partial_tracking_pct=0.15))
    assert "⚠ 15% pages w/o project_name" in mk


def test_marker_orphan_no_project_name():
    mk = D.mvp_cell_markers(_score("__orphan_deadlink__", "MISSING_PROJECT_NAME"))
    assert "⚠ no project_name" in mk


def test_marker_ga_only_posthog_blind():
    mk = D.mvp_cell_markers(_score("x", "GA_NO_PH_TRACKING", ga_only=True, metrics={"ga_clicks": 103}))
    assert "⚠ 103 paid clicks, PostHog blind" in mk


def test_marker_stalled():
    mk = D.mvp_cell_markers(_score("x", "INSUFFICIENT_DATA",
                                   metrics={"stalled_bucket": "stalled"}))
    assert "🧟 stalled" in mk


def test_marker_stalled_slow():
    mk = D.mvp_cell_markers(_score("x", "INSUFFICIENT_DATA",
                                   metrics={"stalled_bucket": "stalled_slow"}))
    assert "🧟 stalled (slow)" in mk


def test_marker_absent_when_stalled_bucket_none():
    mk = D.mvp_cell_markers(_score("x", "INSUFFICIENT_DATA",
                                   metrics={"stalled_bucket": "none"}))
    assert not any("stalled" in m for m in mk)


def test_markers_compose():
    s = _score("x", "NO_GO", db_unmapped_reason="project_deleted", partial_tracking_pct=0.04)
    mk = D.mvp_cell_markers(s)
    assert "⚠ DB deleted" in mk and "⚠ 4% pages w/o project_name" in mk


def test_display_name_strips_orphan_sentinel():
    assert D.display_name(_score("__orphan_momrealm__", "MISSING_PROJECT_NAME")) == "momrealm"


# ---------- no_go_reason (4 branches) ----------

def test_reason_deleted():
    r = D.no_go_reason(_score("x", "NO_GO", db_unmapped_reason="project_deleted"), TH)
    assert "Backend deleted" in r


def test_reason_low_conversion():
    s = _score("x", "NO_GO", metrics={"ga_clicks": 200, "true_conv_rate": 0.007})
    assert D.no_go_reason(s, TH) == "Conversion 0.7% < 6%"


def test_reason_cpc_unit_economics():
    s = _score("x", "NO_GO", metrics={
        "ga_clicks": 100, "true_conv_rate": 0.07,
        "cpc_unit_economics_fail": True, "ga_cpc_usd": 3.16,
        "implied_cac_usd": 63.0, "monthly_price_usd": 19.0,
    })
    r = D.no_go_reason(s, TH)
    assert "CPC $3.16 > $2.50 cap" in r and "CAC $63 > $19/mo" in r
    assert "passed 6% conv" in r  # cleared conversion but bid too high


def test_reason_cpc_plus_low_conv():
    s = _score("x", "NO_GO", metrics={
        "ga_clicks": 100, "true_conv_rate": 0.02,
        "cpc_unit_economics_fail": True, "ga_cpc_usd": 4.80,
        "implied_cac_usd": 96.0, "monthly_price_usd": 49.0,
    })
    r = D.no_go_reason(s, TH)
    assert r.startswith("Conversion 2.0% < 6%  +  CPC")


def test_reason_under_floor():
    s = _score("x", "NO_GO", metrics={"ga_clicks": 40, "true_conv_rate": 0.0})
    assert D.no_go_reason(s, TH) == "Under 100-visitor floor"


# ---------- verdict_summary_counts ----------

def test_summary_counts():
    scores = [
        _score("a", "GO", metrics={"db_signups_real": 1}),
        _score("b", "NO_GO", metrics={"ga_clicks": 1}),
        _score("c", "NO_GO", metrics={"ga_clicks": 2}),
        _score("d", "INSUFFICIENT_DATA"),
        _score("e", "MISSING_PROJECT_NAME"),
    ]
    assert D.verdict_summary_counts(scores) == {"go": 1, "no_go": 2, "insuf": 1, "fix": 1, "promoted": 0}


# ---------- emit_docx: graceful skip + happy path ----------

def test_emit_docx_graceful_skip_when_absent():
    scores = [_score("a", "GO", metrics={"db_signups_real": 1, "ga_clicks": 100})]
    with patch.object(D, "_import_docx", return_value=None):
        with tempfile.TemporaryDirectory() as td:
            out = os.path.join(td, "report.docx")
            ok, msg = D.emit_docx(scores, TH, out)
            assert ok is False
            assert "python-docx not installed" in msg
            assert not os.path.exists(out)  # wrote nothing, did not raise


def test_emit_docx_dry_run_no_file():
    scores = [_score("a", "GO", metrics={"db_signups_real": 1})]
    with tempfile.TemporaryDirectory() as td:
        out = os.path.join(td, "report.docx")
        ok, msg = D.emit_docx(scores, TH, out, dry_run=True)
        # ok depends on python-docx presence; either way no file is written in dry-run
        assert not os.path.exists(out)
        if ok:
            assert "DRY-RUN" in msg


def test_emit_docx_happy_path():
    docx = D._import_docx()
    if docx is None:
        return  # python-docx not installed in this env; happy path covered in CI
    from docx import Document
    scores = [
        _score("muse", "GO", metrics={"ga_clicks": 134, "ph_signups": 9, "db_signups_real": 21,
                                      "true_conv_rate": 0.157, "signup_source": "db_real"},
               signup_events=["signup_complete"]),
        _score("x-predict", "NO_GO", lifecycle_status="killed",
               db_unmapped_reason="project_deleted", metrics={"ga_clicks": 2055, "money_leak": True}),
        _score("dvara", "NO_GO", metrics={"ga_clicks": 112, "ph_signups": 4, "db_signups_real": 4,
                                          "true_conv_rate": 0.036, "signup_source": "db_real"}),
        _score("kansei", "INSUFFICIENT_DATA", metrics={"ga_clicks": 91, "db_signups_real": 3}),
        _score("__orphan_deadlink__", "MISSING_PROJECT_NAME", metrics={}),
    ]
    with tempfile.TemporaryDirectory() as td:
        out = os.path.join(td, "report.docx")
        ok, msg = D.emit_docx(scores, TH, out, gen_date="2026-06-29")
        assert ok is True
        assert os.path.exists(out)
        doc = Document(out)
        # summary mini-table + GO + NO_GO + INSUF + FIX
        assert len(doc.tables) == 5
        assert doc.tables[1].rows[0].cells[-1].text == "SIGNUP EVENTS"  # GO table
        assert doc.tables[2].rows[0].cells[-1].text == "WHY NO_GO"      # NO_GO table


if __name__ == "__main__":
    import inspect

    failed = passed = 0
    for name, fn in list(globals().items()):
        if name.startswith("test_") and callable(fn) and not inspect.signature(fn).parameters:
            try:
                fn()
                print(f"PASS  {name}")
                passed += 1
            except Exception as e:  # noqa: BLE001
                print(f"FAIL  {name}: {e!r}")
                failed += 1
    print(f"\n{passed} passed, {failed} failed")
    sys.exit(1 if failed else 0)
