#!/usr/bin/env python3
"""Render the Phase-1 cross-MVP decision report as a .docx (state x4, best-effort).

Reached via iterate_cross_verdicts.py --emit-docx. Sibling of the --emit-team-message
stdout hand-off: the verdict module owns the flag and calls emit_docx() here, so the
optional python-docx dependency is isolated to this one file. When python-docx is
absent the pipeline must NOT break — emit_docx returns (False, msg), writes nothing,
and the run continues (stdout table + team message are produced regardless).

The report buckets every scored MVP into GO / NO_GO / INSUF / FIX TRACKING and
renders the format validated with the team (4 tables; NO_GO carries a WHY column
instead of SIGNUP EVENTS, GO carries PAID %). All inputs come from the score record
(iterate_cross_verdicts.compute_headline_verdict) — no extra joins needed.

The bucket/marker/reason logic is pure and unit-testable without python-docx.
"""
from __future__ import annotations

DEFAULT_OUTPUT = ".runs/iterate-cross-phase1-decision.docx"

# Verdict enum (kept local to avoid importing the heavy verdicts module).
GO = "GO"
NO_GO = "NO_GO"
INSUFFICIENT = "INSUFFICIENT_DATA"
NO_DATA = "NO_DATA"
MISSING_PROJECT_NAME = "MISSING_PROJECT_NAME"
GA_NO_PH_TRACKING = "GA_NO_PH_TRACKING"

DEFAULT_MAX_CPC = 2.5


# ---------- optional-dependency boundary ----------

def _import_docx():
    """Return the python-docx module, or None when it is not installed.

    The ONLY import site for python-docx. Keeping it behind a function means
    importing this module never fails, so the pure helpers below stay usable
    (and unit-testable) on machines without the dependency.
    """
    try:
        import docx  # noqa: F401
        return docx
    except ImportError:
        return None


# ---------- pure helpers (no python-docx, no I/O) ----------

def _m(score: dict) -> dict:
    return score.get("metrics") or {}


def _ga(score: dict) -> int:
    return _m(score).get("ga_clicks") or 0


def _db_real(score: dict):
    return _m(score).get("db_signups_real")


def _ph(score: dict):
    return _m(score).get("ph_signups")


def _ga_phase2(score: dict) -> int:
    return _m(score).get("ga_clicks_phase2") or 0


def _visitors(score: dict) -> int:
    """Verdict denominator: phase1-scoped GA clicks (blended minus phase2
    campaigns) when GA data is present, else PostHog gclid visitors. Mirrors
    compute_headline_verdict's denominator so WHY strings match the verdict math.
    Defensive: records predating the split lack ga_clicks_phase2 → phase2=0."""
    ga = _ga(score)
    if ga:
        return max(ga - _ga_phase2(score), 0)
    return _m(score).get("gclid_visitors") or 0


def is_orphan(score: dict) -> bool:
    name = score.get("name") or ""
    return (
        score.get("headline_verdict") == MISSING_PROJECT_NAME
        or name.startswith("__orphan_")
        or score.get("db_unmapped_reason") == "orphan"
    )


def is_deleted(score: dict) -> bool:
    return (
        score.get("db_unmapped_reason") in ("project_deleted", "archived_killed")
        or score.get("lifecycle_status") == "killed"
    )


# Conversion-class verdicts eligible for the promoted bucket (data-integrity
# verdicts stay in FIX — a promoted MVP with broken tracking belongs on the
# fix worklist).
_PROMOTED_BUCKET_VERDICTS = (GO, NO_GO, INSUFFICIENT)


def is_promoted(score: dict) -> bool:
    """Promoted partition rule — local twin of
    iterate_cross_verdicts._score_is_promoted (this module deliberately avoids
    importing the heavy verdicts module; the partition-parity test pins the two
    implementations together). Precedence: dead (deleted/killed/money_leak)
    beats promoted."""
    return (
        score.get("lifecycle_status") == "promoted"
        and score.get("headline_verdict") in _PROMOTED_BUCKET_VERDICTS
        and not _m(score).get("money_leak")
        and not is_deleted(score)
    )


def display_name(score: dict) -> str:
    n = score.get("name") or ""
    if n.startswith("__orphan_") and n.endswith("__"):
        return n[len("__orphan_"):-2] or "(no host)"
    return n


def mvp_cell_markers(score: dict) -> list[str]:
    """Warning suffixes appended to the MVP cell (order = severity of trust loss)."""
    mk: list[str] = []
    if is_orphan(score):
        mk.append("⚠ no project_name")
    if score.get("ga_only"):
        mk.append(f"⚠ {_ga(score)} paid clicks, PostHog blind")
    if is_deleted(score):
        mk.append("⚠ DB deleted")
    if score.get("db_unmapped_reason") in ("no_match", "no_match_neither") and not is_orphan(score) and not score.get("ga_only"):
        mk.append("⚠ no DB access")
    pt = score.get("partial_tracking_pct")
    if pt and pt > 0:
        mk.append(f"⚠ {round(pt * 100)}% pages w/o project_name")
    # Stalled-flow triage (state-x3 annotate_stalled) — only ever set on INSUF
    # rows, so this surfaces in the ⏳ Insuf. table without any shape change.
    sb = _m(score).get("stalled_bucket")
    if sb in ("stalled", "stalled_slow"):
        mk.append("🧟 stalled" + (" (slow)" if sb == "stalled_slow" else ""))
    return mk


def name_cell(score: dict) -> str:
    mk = mvp_cell_markers(score)
    return display_name(score) + ("  " + "  ".join(mk) if mk else "")


def no_go_reason(score: dict, thresholds: dict) -> str:
    """Why a NO_GO row was stopped (the WHY column). Precedence:
    deleted backend > CPC unit-economics > low conversion > under floor.
    """
    m = _m(score)
    if is_deleted(score):
        return "Backend deleted / retired — no ground truth"
    floor = (thresholds or {}).get("visitors_floor", 100)
    conv_go = (thresholds or {}).get("conv_rate_go", 0.06)
    v = _visitors(score)
    conv = m.get("true_conv_rate") or 0
    low_conv = v >= floor and conv < conv_go
    if m.get("cpc_unit_economics_fail"):
        cpc = m.get("ga_cpc_usd")
        cac = m.get("implied_cac_usd")
        price = m.get("monthly_price_usd")
        cap = (thresholds or {}).get("max_cpc", DEFAULT_MAX_CPC)
        txt = f"CPC ${cpc:.2f} > ${cap:.2f} cap" if cpc is not None else "CPC over cap"
        if cac is not None and price is not None:
            txt += f" → CAC ${cac:.0f} > ${price:.0f}/mo"
        if low_conv:
            return f"Conversion {conv * 100:.1f}% < 6%  +  {txt}"
        if v >= floor and conv >= conv_go:
            txt += "  (passed 6% conv — bid too high)"
        return txt
    if low_conv:
        return f"Conversion {conv * 100:.1f}% < 6%"
    if v < floor:
        return f"Under {floor}-visitor floor"
    return "—"


def bucket_scores(scores: list[dict]) -> dict:
    """Partition scores into the 5 report buckets.

    Precedence FIX > dead-absorption (killed/deleted/money-leak → no_go) >
    PROMOTED (any conversion-class reference verdict, incl. NO_GO) > plain
    verdict NO_GO > GO > INSUF, via a `seen` guard so each MVP lands once
    (a killed GO, a promoted NO_GO, an orphan, etc. cannot appear twice).
    """
    go, no_go, insuf, fix, promoted = [], [], [], [], []
    seen: set[int] = set()

    def take(score):
        sid = id(score)
        if sid in seen:
            return True
        seen.add(sid)
        return False

    for s in scores:
        v = s.get("headline_verdict")
        if v in (MISSING_PROJECT_NAME, GA_NO_PH_TRACKING, NO_DATA):
            if not take(s):
                fix.append(s)
    # Dead absorption runs BEFORE the promoted pass so a promoted MVP whose
    # backend later dies renders as a stop, not as "promoted, no action".
    for s in scores:
        if s.get("lifecycle_status") == "killed" or _m(s).get("money_leak") or s.get("db_unmapped_reason") in ("project_deleted", "archived_killed"):
            if not take(s):
                no_go.append(s)
    # Promoted pass must run BEFORE the plain-NO_GO pass: a promoted MVP's
    # phase-1 reference verdict can legitimately be NO_GO (phase1 flight wound
    # down) without it belonging in the stop list.
    for s in scores:
        if is_promoted(s):
            if not take(s):
                promoted.append(s)
    for s in scores:
        if s.get("headline_verdict") == NO_GO:
            if not take(s):
                no_go.append(s)
    for s in scores:
        if s.get("headline_verdict") == GO:
            if not take(s):
                go.append(s)
    for s in scores:
        if s.get("headline_verdict") == INSUFFICIENT:
            if not take(s):
                insuf.append(s)

    go.sort(key=lambda s: -(_db_real(s) or 0))
    no_go.sort(key=lambda s: -_ga(s))
    insuf.sort(key=lambda s: -_ga(s))
    fix.sort(key=lambda s: -_ga(s))
    promoted.sort(key=lambda s: -_ga_phase2(s))
    return {"go": go, "no_go": no_go, "insuf": insuf, "fix": fix, "promoted": promoted}


def verdict_summary_counts(scores: list[dict]) -> dict:
    b = bucket_scores(scores)
    return {
        "go": len(b["go"]),
        "no_go": len(b["no_go"]),
        "insuf": len(b["insuf"]),
        "fix": len(b["fix"]),
        "promoted": len(b["promoted"]),
    }


def _fmt_int(v) -> str:
    return "—" if not v else f"{v:,}"


def _fmt_sig(v) -> str:
    return "—" if v is None else str(v)


def _fmt_conv(score: dict) -> str:
    return f"{(_m(score).get('true_conv_rate') or 0) * 100:.1f}%"


def _fmt_paid(score: dict) -> str:
    g = _ga(score)
    p = _ph(score) or 0
    return f"{(p / g * 100):.1f}%" if g else "0.0%"


def _fmt_source(score: dict) -> str:
    src = _m(score).get("signup_source")
    if src in ("db_real", "db_paid", "db_real_zero"):
        return "DB"
    if src == "ph":
        return "PH"
    return "—"


def _fmt_events(score: dict) -> str:
    ev = score.get("signup_events") or []
    return ", ".join(ev) if ev else "—"


def _aggregates(scores: list[dict]) -> dict:
    on_ga = sum(1 for s in scores if _ga(s) > 0)
    return {
        "n": len(scores),
        "on_google_ads": on_ga,
        "sum_ga": sum(_ga(s) for s in scores),
        "sum_ph": sum(_ph(s) or 0 for s in scores),
        "sum_db": sum(_db_real(s) or 0 for s in scores),
        "sum_eff": sum(_m(s).get("effective_signups") or 0 for s in scores),
        "n_deleted": sum(1 for s in scores if is_deleted(s)),
        "no_db_access": [
            display_name(s) for s in scores
            if s.get("db_unmapped_reason") in ("no_match", "no_match_neither")
            and not is_orphan(s) and not s.get("ga_only")
        ],
    }


# ---------- docx writer (the only part needing python-docx) ----------

def emit_docx(scores: list[dict], thresholds: dict, out_path: str, dry_run: bool = False, gen_date: str | None = None):
    """Build the decision report .docx at out_path.

    Returns (ok: bool, message: str). When python-docx is absent, returns
    (False, "...skipping...") and writes nothing — the caller treats this as a
    best-effort skip, never an error.
    """
    docx = _import_docx()
    if docx is None:
        return False, "python-docx not installed; skipping .docx report (pip install python-docx to enable)"
    if dry_run:
        b = bucket_scores(scores)
        return True, (
            f"DRY-RUN: would write {out_path} "
            f"(GO={len(b['go'])} NO_GO={len(b['no_go'])} INSUF={len(b['insuf'])} "
            f"FIX={len(b['fix'])} PROMOTED={len(b['promoted'])})"
        )

    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    thresholds = thresholds or {}
    floor = thresholds.get("visitors_floor", 100)
    cap = thresholds.get("max_cpc", DEFAULT_MAX_CPC)
    b = bucket_scores(scores)
    agg = _aggregates(scores)

    NAVY = RGBColor(0x1F, 0x2A, 0x44)
    GREEN = RGBColor(0x16, 0x7A, 0x3B)
    RED = RGBColor(0xB0, 0x2A, 0x2A)
    AMBER = RGBColor(0x9A, 0x6A, 0x00)
    GREY = RGBColor(0x55, 0x55, 0x55)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    def para(text, size=10, bold=False, color=None, after=6, before=0, align=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(after)
        p.paragraph_format.space_before = Pt(before)
        if align:
            p.alignment = align
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        if color is not None:
            r.font.color.rgb = color
        return p

    def shade(cell, hexcolor):
        tcPr = cell._tc.get_or_add_tcPr()
        sh = OxmlElement("w:shd")
        sh.set(qn("w:val"), "clear")
        sh.set(qn("w:color"), "auto")
        sh.set(qn("w:fill"), hexcolor)
        tcPr.append(sh)

    def set_cell(cell, text, bold=False, size=8.5, color=None, align=None):
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)
        if align:
            p.alignment = align
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        if color is not None:
            r.font.color.rgb = color

    def build_table(rows, headers, widths, vcolor):
        t = doc.add_table(rows=1, cols=len(headers))
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        t.style = "Table Grid"
        for j, h in enumerate(headers):
            set_cell(t.rows[0].cells[j], h, bold=True, size=8.5, color=WHITE)
            shade(t.rows[0].cells[j], "1F2A44")
        for row in rows:
            cells = t.add_row().cells
            for j, val in enumerate(row):
                set_cell(cells[j], val, size=8.5)
            set_cell(cells[0], row[0], bold=True, size=8.5, color=vcolor)
        for j, w in enumerate(widths):
            for r in t.rows:
                r.cells[j].width = Inches(w)
        return t

    # Header
    para("Phase 1 — Cross-MVP Final Decision Review", size=17, bold=True, color=NAVY, after=2)
    date_str = f" · generated {gen_date}" if gen_date else ""
    para(
        f"{agg['n']} MVPs ({agg['on_google_ads']} on Google Ads) · floor {floor} visitors · "
        f"GO threshold 6% conversion · CPC cap ${cap:.2f}{date_str}",
        size=9, color=GREY, after=10,
    )

    para("What we need you to verify", size=12, bold=True, color=NAVY, after=2)
    para(
        "Please sanity-check the two ground-truth columns on the GO and NO_GO rows: GA CLICKS (from the "
        "Google Ads MCC export) and DB SIGNUPS (real, after test/team filtering). There is no sign-off "
        "column — if a number looks wrong to you, just leave a comment on that row (or ping me) and I'll "
        "investigate before we act on it. The NO_GO table carries a WHY column so you can see whether each "
        "stop is driven by low conversion, the CPC cap, or a deleted backend. The GO table shows PAID % — "
        "conversion using only gclid-attributed PostHog signups — to expose deploys whose headline "
        "conversion leans on organic traffic.",
        size=9.5, after=8,
    )

    note = (
        f"Note — {agg['n_deleted']} deleted/retired-backend MVPs show ⚠ DB deleted (Supabase torn down or "
        f"MVP killed; historical ground truth unrecoverable → STOP and archived)."
    )
    if agg["no_db_access"]:
        note += (
            f" {', '.join(agg['no_db_access'])} show ⚠ no DB access (no Supabase/Railway match — verdict "
            f"falls back to PostHog)."
        )
    para(note, size=9, color=GREY, after=10)

    # Verdict summary mini-table
    para("Verdict summary", size=12, bold=True, color=NAVY, after=4)
    st = doc.add_table(rows=1, cols=4)
    st.style = "Table Grid"
    st.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, (txt, fill) in enumerate([
        (f"PROMOTE (GO)  {len(b['go'])}", "16A34A"),
        (f"STOP (NO_GO)  {len(b['no_go'])}", "B02A2A"),
        (f"KEEP (INSUF.)  {len(b['insuf'])}", "9A6A00"),
        (f"IN PHASE 2  {len(b['promoted'])}", "1F2A44"),
    ]):
        set_cell(st.rows[0].cells[j], txt, bold=True, size=11, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(st.rows[0].cells[j], fill)
        st.rows[0].cells[j].width = Inches(1.75)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    para(
        f"Σ GA clicks: {agg['sum_ga']:,}    ·    Σ PH signups: {agg['sum_ph']}    ·    "
        f"Σ DB signups (real, post-filter): {agg['sum_db']}    ·    Σ effective signups: {agg['sum_eff']}"
        f"    ·    {len(b['fix'])} deploys need tracking fixes (below)",
        size=9, bold=True, color=NAVY, after=12,
    )

    # GO table (has PAID %)
    para("✅ GO — promote to Phase 2", size=12, bold=True, color=GREEN, after=4)
    build_table(
        [["✅ GO", name_cell(s), _fmt_int(_ga(s)), _fmt_sig(_ph(s)), _fmt_sig(_db_real(s)),
          _fmt_conv(s), _fmt_paid(s), _fmt_source(s), _fmt_events(s)] for s in b["go"]],
        ["VERDICT", "MVP", "GA CLICKS", "PH SIGNUPS", "DB SIGNUPS", "CONV %", "PAID %", "SOURCE", "SIGNUP EVENTS"],
        [0.7, 1.9, 0.7, 0.7, 0.7, 0.6, 0.6, 0.5, 1.8], GREEN,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Promoted table (operator-confirmed lifecycle_status=promoted; excluded
    # from the decision buckets above — Phase 1 does not re-litigate them).
    if b["promoted"]:
        para("🚀 Promoted to Phase 2 — no Phase 1 action", size=12, bold=True, color=NAVY, after=4)
        build_table(
            [["🚀 IN P2", name_cell(s), (s.get("lifecycle_status_at") or "—")[:10],
              s.get("headline_verdict") or "—", _fmt_int(_ga_phase2(s)), _fmt_conv(s)]
             for s in b["promoted"]],
            ["STATUS", "MVP", "PROMOTED AT", "PHASE-1 REF VERDICT", "PHASE-2 CLICKS", "P1 CONV %"],
            [0.8, 2.2, 1.0, 1.4, 1.0, 0.8], NAVY,
        )
        para(
            "Reference verdicts only — the operative call for these MVPs is the pay-intent verdict "
            "from /iterate --cross --phase2.",
            size=9, color=GREY, after=8, before=2,
        )
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # NO_GO table (WHY column replaces SIGNUP EVENTS)
    para("❌ NO_GO — stop spend", size=12, bold=True, color=RED, after=4)
    build_table(
        [["❌ NO_GO", name_cell(s), _fmt_int(_ga(s)), _fmt_sig(_ph(s)), _fmt_sig(_db_real(s)),
          _fmt_conv(s), _fmt_source(s), no_go_reason(s, thresholds)] for s in b["no_go"]],
        ["VERDICT", "MVP", "GA CLICKS", "PH SIGNUPS", "DB SIGNUPS", "CONV %", "SOURCE", "WHY NO_GO"],
        [0.75, 2.05, 0.62, 0.62, 0.62, 0.55, 0.45, 2.34], RED,
    )
    n_conv = sum(1 for s in b["no_go"] if not is_deleted(s) and not _m(s).get("cpc_unit_economics_fail"))
    n_cpc = sum(1 for s in b["no_go"] if not is_deleted(s) and _m(s).get("cpc_unit_economics_fail"))
    n_del = sum(1 for s in b["no_go"] if is_deleted(s))
    para(
        f"Why these stopped — {n_conv} fail on conversion (≥{floor} paid visitors, under the 6% bar), "
        f"{n_cpc} fail on CPC unit-economics (avg CPC over the ${cap:.2f} cap → implied CAC exceeds one "
        f"month's price), and {n_del} have a deleted/retired backend (no ground truth → archived). The CPC "
        "stops are revivable: lower the campaign's max-CPC bid to the cap, or approve an exception, then re-run.",
        size=9, color=GREY, after=12, before=4,
    )

    # INSUF table
    para("⏳ Insuf. — keep running until floor reached", size=12, bold=True, color=AMBER, after=4)
    build_table(
        [["⏳ Insuf.", name_cell(s), _fmt_int(_ga(s)), _fmt_sig(_ph(s)), _fmt_sig(_db_real(s)),
          _fmt_conv(s), _fmt_source(s), _fmt_events(s)] for s in b["insuf"]],
        ["VERDICT", "MVP", "GA CLICKS", "PH SIGNUPS", "DB SIGNUPS", "CONV %", "SOURCE", "SIGNUP EVENTS"],
        [0.75, 2.4, 0.7, 0.7, 0.7, 0.6, 0.5, 1.8], AMBER,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # FIX TRACKING table
    para(
        "🚨 FIX TRACKING — verdict blocked; fix analytics.ts PROJECT_NAME, then re-run /iterate",
        size=12, bold=True, color=NAVY, after=4,
    )
    build_table(
        [["🚨 FIX", name_cell(s), _fmt_int(_ga(s)), _fmt_sig(_ph(s)), _fmt_sig(_db_real(s)),
          _fmt_conv(s), _fmt_source(s), _fmt_events(s)] for s in b["fix"]],
        ["VERDICT", "MVP", "GA CLICKS", "PH SIGNUPS", "DB SIGNUPS", "CONV %", "SOURCE", "SIGNUP EVENTS"],
        [0.7, 2.9, 0.6, 0.7, 0.7, 0.6, 0.5, 1.4], NAVY,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Methodology
    para("Methodology", size=12, bold=True, color=NAVY, after=4)
    for mtext in [
        "GA clicks come from the Google Ads MCC CSV export. PH = PostHog. DB signups come from each MVP's "
        "Supabase or Railway database, counting real rows after removing test and team accounts via the email filter.",
        "Visitors (verdict denominator) = GA clicks when present, else PostHog paid-traffic visitors. CONV % = "
        "effective signups ÷ visitors. Effective signups = trusted DB real count first, PostHog paid signups as fallback.",
        "WHY NO_GO column: (1) Conversion X% < 6% — enough paid traffic but below the GO bar; (2) CPC over the "
        f"${cap:.2f} cap → CAC > monthly price — cannot pay back at this click price (fires even before the "
        f"{floor}-visitor floor, so a high-CPC deploy is stopped without burning more spend); (3) Backend "
        "deleted/retired — ground truth is gone, archived.",
        "PAID % (GO table only) = PostHog gclid-attributed signups ÷ GA clicks — a conservative paid-only "
        "conversion. CONV % can exceed PAID % when a DB count includes organic signups (the DB table has no "
        "gclid column). Trust PAID % for the promote decision.",
        "Source column: DB = trusted database ground truth used. PH = PostHog count used (DB unavailable). — = "
        "no signal / blocked.",
        "FIX TRACKING bucket: deploys whose verdict cannot be trusted — Google Ads is spending but PostHog "
        "records zero events (analytics.ts missing or PROJECT_NAME mismatch), or events stream with no project_name.",
    ]:
        para(mtext, size=9, color=GREY, after=4)

    doc.save(out_path)
    return True, f"Wrote {out_path}"
